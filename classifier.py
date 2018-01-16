#!usr/bin/python

import datetime
from email.mime.text import MIMEText
from email.generator import Generator
from email.header import Header
import html
import os
import os.path as osp
from pathlib import Path
import string

import pandas as pd
import numpy as np

from nltk.corpus import stopwords
from nltk.stem.snowball import SnowballStemmer

from sklearn.linear_model import LogisticRegression
from sklearn.model_selection import (GridSearchCV, StratifiedKFold,
                                     train_test_split)
from sklearn.metrics import (f1_score, make_scorer, confusion_matrix,
                             classification_report)
from sklearn.externals import joblib
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.pipeline import Pipeline

import docx
from docx.shared import Inches
from hyperlink import add_hyperlink

# https://pypi.python.org/pypi/emoji/
import emoji


"""
Description

Keyword arguments:
:param X:            type - description

:returns:               type        -- description
"""


TODAY = datetime.date.today()


class TweetClassifier(object):
    """
    Tweet_Classifier classifies tweets into interesting and not interesting
    based on a Logistic Regression model.

    It takes in a Pandas DataFrame pre-downloaded by TweetCollector and extends
    the DataFrame by the classification and the probability of this prediction.

    :param data:            Pandas DataFrame
                            Columns contain:
                            since_id    -- string of numbers corresponding
                                           to tweet status id
                            created_at  -- string in date format of posting
                                           time, e.g. 2017-12-06 16:37:49
                            tweet       -- string of tweet
                            feed        -- string of name of twitter feed

    :param train:           Pandas DataFrame
                            Columns contain:
                            since_id    -- string of numbers corresponding
                                           to tweet status id
                            created_at  -- string in date format of posting
                                           time, e.g. 2017-12-06 16:37:49
                            tweet       -- string of tweet
                            feed        -- string of name of twitter feed
                            interesting -- int signifying interesting or not

    :param ignore saved:    Boolean - flag: ignores saved classifier and re-
                            calibrates if true (default:False)

    Methods defined here

    def create_classifier(self, X=None, y=None, ignore_saved_model=False)
        Returns a fully trained Logistic Regression model ready to be used
        for predictions.
        Exception: In case there are no pickled classifier, no hyperparameters
        and no training data supplied, None is returned.

    def calibrate(self, X, y, full_calibration=False)
        Calibrate a prediction model based on a sklearn.Pipeline of a
        TfidfVectorizer using self.text_process as the analyzer and a Logistic
        Regression model.

    def predict(self, pref_prob=1)
        Uses self.clf's classification model to predict if a Pandas
        Series of Tweets are interesting (numpy.array[,]) and the
        respective probabilites (numpy.array[[],]).

    def check_accuracy(self, X_text, y_test)
        Prints Confusion Matrix and Classification Report for the test split
        after model has been fit.

    def save_classifier(self)
        Saves classification model and the respective hyperparameters as pickle
        file each using sklearn.joblib.dump.

    def save_as_txt(self, fname='report{}.txt'.format(TODAY), cut_off=0.25)
        Saves the classification outcome as a text file

    def save_as_doc(self, fname='report{}.docx'.format(TODAY),
                    doc_title='Tweet Report {}'.format(TODAY), cut_off=0.25)
        Saves the classification outcome in a table of a docx file.

    def text_process(self, mess)
        Converts string into list of bag of words.
    """

    def __init__(self, data, train=None, ignore_saved=False,
                 full_calibration=False):

        # Paths
        # path to data directory
        self._dest = osp.join(os.getcwd(), 'data')
        # path to pickled classifier
        self._clf_path = osp.join(self._dest, 'classifier.pkl')
        # path to pre-saved hyperparams
        self._hyperparams = osp.join(self._dest, 'hyperparams.pkl')

        # information fed into class
        self.data = data
        self.train_set = train
        self.X_train, self.X_test, self.y_train, self.y_test = \
            train_test_split(train['tweet'], train['interesting'],
                             test_size=0.33, random_state=42)

        self.full_cal = full_calibration

        # Adjust to a property function based on ending either csv or xlsx
        self.tfidf = TfidfVectorizer(analyzer=self.text_process,
                                     ngram_range=(1, 1))

        self.clf = self.create_classifier(self.X_train,
                                          self.y_train,
                                          ignore_saved_model=ignore_saved)

        # Pandas Series or None
        self.clf_text = self.data['tweet']
        # Pandas series or None
        self.prediction = None
        # Pandas series or None
        self.proba = None
        # Pandas Dataframe of all info or None
        self.assembled = None

    def create_classifier(self, X=None, y=None, ignore_saved_model=False):
        """
        Returns a fully trained Logistic Regression model ready to be used
        for predictions.
        Exception: In case there are no pickled classifier, no hyperparameters
        and no training data supplied, None is returned.

        Keyword arguments:
        :param X:                   Pandas Series - Collection of Tweets
        :param y:                   Pandas Series - Collection of Integers
                                                    1 = Interesting, 0 = Not
        :param ignore_saved_model:  Boolean       - flag: Ignores pickled model
                                                    Allows for recalibration if
                                                    new training data given.

        :returns:                   trained sklearn.Pipeline
                                        ('vect', tfidf),
                                        ('clf', LogisticRegression())
        """

        if Path(self._clf_path).is_file() and not ignore_saved_model:
            clf = joblib.load(open(self._clf_path, 'rb'))
            print('Classifier successfully loaded')
            return clf

        else:
            print('No classifier pre-saved. Please wait while a new ',
                  'classifier is being calibrated...')

            if X is not None and y is not None:
                clf = self.calibrate(X, y, full_calibration=self.full_cal)
                print('Classifier has been calibrated.')
                return clf

            else:
                print('You have not defined any training data.',
                      ' Please define:\n')
                if X is None:
                    print('X - pandas Dataframe defining data features.\n')
                if y is None:
                    print('y - pandas Data Series defining the expected',
                          ' classification.\n')
                return None

    def calibrate(self, X, y, full_calibration=False):
        """
        Calibrate a prediction model based on a sklearn.Pipeline of a
        TfidfVectorizer using self.text_process as the analyzer and a Logistic
        Regression model.
        The training data allow to classify tweets downloaded from pre-defined
        feeds.

        Keyword arguments:
        :param X:                   Pandas Series - Collection of Tweets
        :param y:                   Pandas Series - Collection of Integers
                                                    1 = Interesting, 0 = Not
        :param full_calibration:    Boolean       - flag: results in full
                                                    calibration of a new model
                                                    using sklearn.GridSearchCV
                                                    (Will take several minutes)

        :returns:                   trained sklearn.Pipeline
                                        ('vect', tfidf),
                                        ('clf', LogisticRegression())
        """

        tfidf = TfidfVectorizer(analyzer=self.text_process)
        pipe = Pipeline([
            ('vect', tfidf),
            ('clf', LogisticRegression())
        ])

        # F1 score is measured for the 1 label, as it is more interesting
        # to have all interesting tweets correct, than some uninteresting
        # ones mislabelled
        scorer = make_scorer(f1_score, pos_label=1)

        clf = None

        if Path(self._hyperparams).is_file() and not full_calibration:
            param_grid = joblib.load(open(self._hyperparams, 'rb'))
            pipe.set_params(**param_grid)
            pipe.fit(X, y)
            clf = pipe

        else:
            param_grid = [
                {'vect__ngram_range': [(1, 1), (2, 2), (1, 2)],
                 'vect__use_idf': [False, True],
                 'clf__penalty': ['l1', 'l2'],
                 'clf__C': [0.1, 1.0, 10.0, 100.0]
                 }
            ]
            grid = GridSearchCV(pipe, param_grid, scoring=scorer,
                                cv=10, verbose=10, n_jobs=-1)
            grid.fit(X, y)
            clf = grid.best_estimator_

        return clf

    def predict(self, pref_prob=1):
        """
        Uses self.clf's classification model to predict if a Pandas
        Series of Tweets are interesting (numpy.array[,]) and the
        respective probabilites (numpy.array[[],]).

        The resulting predictions are merged together with originally
        passed self.data on the tweets as self.assembled.

        Keyword arguments:
        :param pref_prob:   int - Passes index of which probability should
                                  be listed in self.assembled. (default=1,
                                  i.e. interesting = 1)

        :returns:           None

        :side-effects:      self.assembled set to Pandas DataFrame
                            Pandas DataFrame with Columns
                            since_id    -- string of numbers corresponding
                                           to tweet status id
                            created_at  -- string in date format of posting
                                           time, e.g. 2017-12-06 16:37:49
                            tweet       -- string of tweet
                            feed        -- string of name of twitter feed
                            interesting -- int signifying interesting or not
                            probability -- numpy.float of probability of
                                           interesting
        """

        self.prediction = self.clf.predict(self.clf_text)

        self.proba = self.clf.predict_proba(self.clf_text)

        a = pd.DataFrame({'tweet': self.clf_text,
                          'interesting': self.prediction,
                          'probability': self.proba[:, pref_prob]})
        self.assembled = pd.merge(self.data, a, how='outer', left_on='tweet',
                                  right_on='tweet')

    def print_accuracy(self, X_test=None, y_test=None):
        """
        Prints Confusion Matrix and Classification Report for the test split
        after model has been fit.

        Keyword arguments:
        :param X_test:  Pandas Series of strings - The downloaded tweets

        :param y_test:  Pandas Series of ints - 0 or 1s (interesting or not)

        :returns:           None

        :side-effects:  prints confusion matrix of prediction vs. actual
                        test data.

        """

        if X_test is None:
            X_test = self.X_test

        if y_test is None:
            y_test = self.y_test

        prediction = self.clf.predict(X_test)

        # print('Best Score: ', self.clf.best_score_)
        print('\n')
        print(confusion_matrix(prediction, y_test))
        print('\n')
        print(classification_report(prediction, y_test))
        return

    def get_prediction(self):
        """
        Return the Pandas DataFrame containing the tweets as well as the
        predictions or None if no prediction has been done.

        :returns:           None or

                            self.assembled set to Pandas DataFrame
                            Pandas DataFrame with Columns
                            since_id    -- string of numbers corresponding
                                           to tweet status id
                            created_at  -- string in date format of posting
                                           time, e.g. 2017-12-06 16:37:49
                            tweet       -- string of tweet
                            feed        -- string of name of twitter feed
                            interesting -- int signifying interesting or not
                            probability -- numpy.float of probability of
                                           interesting
        """

        return self.assembled

    def save_classifier(self):
        """
        Saves classification model and the respective hyperparameters as pickle
        file each using sklearn.joblib.dump.

        Keyword arguments:
        No params

        :returns:       None

        :side-effects:  Saves a serialized python object structure
                        (pickled file)
        """

        if self.clf is None:
            print('Cannot save NoneType. Are you sure a classifier is loaded?')
        else:
            if not osp.exists(self._dest):
                os.makedirs(self._dest)
            # Classifier saved
            joblib.dump(self.clf, open(self._clf_path, 'wb'),
                        protocol=4)
            print('Classifier saved.')

            # Hyperparameters saved
            joblib.dump(self.clf.get_params(), open(self._hyperparams, 'wb'),
                        protocol=4)
            print('Hyperparameters of Classifier saved.')

    def save_as_txt(self, fname='report{}.txt'.format(TODAY), cut_off=0.25):
        """
        Saves the classification outcome as a text file.

        Columns:
        - Relevance as a percentage
        - Tweet
        - URL of tweet
        - Date in format YYYY-MM-DD HH:MM:SS

        Keyword arguments:
        :param fname:   string - filename or path to filename
                        (default: reportYYYY-MM-DD.txt)
        :param cut_off: float - percentage under which the tweet will be
                        disregarded from the report (default: 0.25)

        :returns:       type        -- description

        :side-effects:  Saves a txt file under fname.
        """

        if self.prediction is None or self.proba is None or \
                self.assembled is None:
            print('No classification has taken place. Please',
                  ' re-use this method after classification has taken place')
            return

        else:
            with open(osp.abspath(fname), 'w', encoding='utf-8') as f:
                f.write('{:12s}\t{}\t{}\t{}\n'.format('Relevance', 'Text',
                                                      'URL', 'Date'))
                f.write('{:12s}\t{}\t{}\t{}\n'.format('=========', '====',
                                                      '===', '===='))
                rel = self.assembled.sort_values(by='probability',
                                                 ascending=False)

                # columns are 'since_id', 'created_at', 'tweet', 'feed',
                # 'interesting', 'probability',
                for since_id, date, tweet, feed, clf, prob in rel.values:
                    if prob >= cut_off:
                        url = 'https://twitter.com/{feed}/status/{since_id}'\
                            .format(feed=feed, since_id=since_id)
                        f.write('{:7.2f}%\t{tweet}\t{url}\t{date}\n'
                                .format(prob * 100,
                                        # HTML Chars and Emojis otherwise not
                                        # correctly shown
                                        tweet=html.unescape(tweet),
                                        url=url, date=date))

            print('Report saved in {}'.format(osp.abspath(fname)))
            return

    def save_as_doc(self, fname='report{}.docx'.format(TODAY),
                    doc_title='Tweet Report {}'.format(TODAY), cut_off=0.25):
        """
        Saves the classification outcome in a table of a docx file.

        Columns:
        - Relevance as a percentage
        - Tweet
        - Feed / Link text feed with hyperlink to URL of tweet
        - Date in format YYYY-MM-DD

        Keyword arguments:
        :param fname:       string - filename or path to filename
                            (default: reportYYYY-MM-DD.docx)
        :param doc_title:   string - title of document
                            (default: Tweet Report YYYY-MM-DD.txt)
        :param cut_off:     float - percentage under which the tweet will be
                            disregarded from the report (default: 0.25)

        :returns:           None

        :side-effects:      Saves a docx file under fname.

        """

        def make_bold(cell):
            """
            Puts text content of cell in bold

            Keyword arguments:
            :param cell:    cell of table

            :returns:       None
            """

            run = cell.paragraphs[0].runs[0]
            font = run.font
            font.bold = True
            return

        def make_hyperlink(cell, url, text, color=None, underline=True):
            """
            Adds a text with hyperlink_url into cell of table.

            Keyword arguments:
            :param cell:        docx.Table._cell of table
            :param url:         string  - url for hyperlink
            :param text:        string  - docx text imbued with url as
                                hyperlink
            :param color:       string  - hex of color
                                (default: None)
            :param underline:   boolean - flag: underline hyperlink

            :returns:       None
            """

            p = cell.paragraphs[0]
            hyperlink = add_hyperlink(p, url, text, None, True)
            return

        def set_col_widths(table, col_widths):
            """
            Sets the widths of columns of table

            table = docx.Table() instance
            col_widths = iterable of ints (width in EMUs)
                         Helper functions like Cm() or Inches() help
            Description

            Keyword arguments:
            :param table:       docx.Table object
            :param col_widths:  list/tuple of int - widths in EMUs
                                Helper functions like Cm() or Inches() help

            :returns:           None
            """

            for col, width in zip(table.columns, col_widths):
                col.width = width
            return

        if self.prediction is None or self.proba is None or \
                self.assembled is None:
            print('No classification has taken place. Please',
                  ' re-use this method after classification has taken place')
            return

        else:
            doc = docx.Document()

            # Create the template
            doc.add_heading(doc_title, 0)

            # Create table
            table = doc.add_table(rows=1, cols=4)
            table.autofit = False

            # Header Cells
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Relevance'
            make_bold(hdr_cells[0])

            hdr_cells[1].text = 'Tweet'
            make_bold(hdr_cells[1])

            hdr_cells[2].text = 'Feed / Link'
            make_bold(hdr_cells[2])

            hdr_cells[3].text = 'Date'
            make_bold(hdr_cells[3])

            # Body of Table
            rel = self.assembled.sort_values(by='probability',
                                             ascending=False)

            # columns are 'since_id', 'created_at', 'tweet', 'feed',
            # 'interesting', 'probability',
            for since_id, date, tweet, feed, clf, prob in rel.values:
                if prob >= cut_off:
                    url = 'https://twitter.com/{feed}/status/{since_id}'\
                        .format(feed=feed, since_id=since_id)
                    row_cells = table.add_row().cells
                    row_cells[0].text = '{:7.2f}%'.format(prob * 100)
                    # HTML Chars and Emojis otherwise not correctly shown
                    row_cells[1].text = html.unescape(tweet)
                    row_cells[2].text = ''
                    make_hyperlink(row_cells[2], url, feed)
                    row_cells[3].text = str(date)[:11]

            set_col_widths(table, (Inches(1.06), Inches(3.0), Inches(0.88),
                                   Inches(1.06)))

            doc.save(osp.abspath(fname))
            print('Report saved in {}'.format(osp.abspath(fname)))
            return

    def save_as_email(self, subject='', sender='', receiver='',
                      fname='report-{}.elm'.format(TODAY), cut_off=0.25):
        """
        Saves the classification outcome as a elm file.

        Columns:
        - Relevance as a percentage
        - Tweet
        - URL of tweet
        - Date in format YYYY-MM-DD HH:MM:SS

        Keyword arguments:
        :param subject: string - Subject of the email
        :param sender:  string - email address of sender
        :param receiver:string or list - recipient address or list of address
        :param fname:   string - filename or path to filename
                        (default: reportYYYY-MM-DD.txt)
        :param cut_off: float - percentage under which the tweet will be
                        disregarded from the report (default: 0.25)

        :returns:       type        -- description

        :side-effects:  Saves a txt file under fname.
        """

        if self.prediction is None or self.proba is None or \
                self.assembled is None:
            print('No classification has taken place. Please',
                  ' re-use this method after classification has taken place')
            return

        else:
            htm = '''<!DOCTYPE html>
                <html>
                <head>
                <style>
                table, th, td {
                    border: 1px solid black;
                }
                </style>
                </head>
                <body>
                <table style="width:100%">
                <tr>
                    <th>Relevance</th>
                    <th>Tweet</th>
                    <th>Feed / Link</th>
                    <th>Date</th>
                </tr>
                '''

            # Body of Table
            rel = self.assembled.sort_values(by='probability',
                                             ascending=False)

            for since_id, date, tweet, feed, clf, prob in rel.values:
                if prob >= cut_off:
                    url = 'https://twitter.com/{feed}/status/{since_id}'\
                          .format(feed=feed, since_id=since_id)
                    addon = '''
                    <tr>
                        <td>{:7.2f}%</td>
                        <td>{}</td>
                        <td><a href="{}">{}</a></td>
                        <td>{}</td>
                    </tr>
                    '''.format(prob * 100, tweet, url, feed, date)
                    htm += addon
            htm += '</table>'

            msg = MIMEText(htm, 'html', 'utf-8')
            msg['Subject'] = Header(subject, 'UTF-8')
            msg['From'] = sender

            # Allow list of multiple addresses
            if type(receiver) is list:
                msg['To'] = ", ".join(receiver)
            else:
                msg['To'] = receiver

            # open a file and save mail to it
            with open(fname, 'w') as out:
                gen = Generator(out)
                gen.flatten(msg)

            print('Report saved in {}'.format(osp.abspath(fname)))
            return

    def text_process(self, mess):
        """
        Converts string into list of bag of words.

        Conversion steps are:
        1. Remove all punctuation
        2. Remove all emojis
        3. Remove all stopwords (based on nltk)
        4. Use stemmer to convert all verbs to base form

        Keyword arguments:
        :param mess:    string

        :returns:       list of strings
        """
        if type(mess) is not str:
            mess = str(mess)

        # Check characters to see if they are in punctuation
        nopunc = [char for char in mess
                  if char not in string.punctuation + "–"]

        # Check if char is an emoji
        noemoji = [char for char in nopunc
                   if char not in emoji.UNICODE_EMOJI.keys()]

        # Join the characters again to form the string.
        nopunc = ''.join(noemoji)

        # Now just remove any stopwords
        no_stop = [word for word in nopunc.split() if word.lower() not in
                   stopwords.words('english') if 'http' not in word.lower()]

        # Stemming
        snowball = SnowballStemmer('english')

        return [snowball.stem(word) for word in no_stop]


if __name__ == '__main__':
    clf = pd.read_excel('example_tweets.xlsx')
    df = pd.read_excel('Training Data.xlsx')
    tc = TweetClassifier(clf, train=df, ignore_saved=True,
                         full_calibration=True)
    tc.print_accuracy()
    # Commented out so that hyperparam.pkl does not always change on commit
    tc.save_classifier()
    tc.predict()
    #tc.save_as_doc()
    #tc.save_as_txt()
    tc.save_as_email('Test', 'bob@work.com', ['tracy@work.com',
                                              'Archibald@work.com'])
